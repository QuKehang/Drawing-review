#!/usr/bin/env python3
"""
本地 RAG 知识库系统 —— 附注判别专用版
基于 Ollama + DeepSeek-R1 + ChromaDB
支持 txt / pdf / docx 文档
"""

import os
import re
import time
from typing import List, Optional
from pathlib import Path

import requests

os.environ.setdefault("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_BASE_URL = os.environ.get("OLLAMA_HOST", "http://localhost:11434")


# ── 自定义 Embedding（直接 HTTP API，不走 ollama Python 包）──

class OllamaHTTPEmbeddings:
    def __init__(self, model: str, base_url: str = "http://localhost:11434"):
        self.model = model
        self.base_url = base_url.rstrip("/")

    def embed_query(self, text: str) -> List[float]:
        return self.embed_documents([text])[0]

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if not texts:
            return []
        all_embeddings = []
        for i in range(0, len(texts), 50):
            batch = texts[i:i + 50]
            for attempt in range(3):
                try:
                    resp = requests.post(
                        f"{self.base_url}/api/embed",
                        json={"model": self.model, "input": batch},
                        timeout=120,
                    )
                    resp.raise_for_status()
                    all_embeddings.extend(resp.json()["embeddings"])
                    break
                except requests.exceptions.RequestException as e:
                    if attempt < 2:
                        time.sleep(2 ** attempt)
                    else:
                        raise RuntimeError(
                            f"Embedding API 调用失败（Ollama 可能未启动）\n"
                            f"请确认:\n"
                            f"  1. Ollama 已运行: ollama serve\n"
                            f"  2. 模型已拉取: ollama pull {self.model}\n"
                            f"  3. API 端口可访问: {self.base_url}/api/embed\n"
                            f"原始错误: {e}"
                        )
        return all_embeddings


from langchain_ollama import OllamaLLM
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate


# ── 文档加载器 ─────────────────────────────────────────────

def _load_txt(file_path: str) -> List[Document]:
    """加载 txt 文件，自动尝试多种编码（UTF-8 / GBK / GB2312）"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return [Document(page_content=f.read(), metadata={"source": file_path})]
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 最终兜底：errors="replace"
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return [Document(page_content=f.read(), metadata={"source": file_path})]


def _load_md(file_path: str) -> List[Document]:
    return _load_txt(file_path)


def _load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader
    docs = []
    for i, page in enumerate(PdfReader(file_path).pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(Document(page_content=text,
                          metadata={"source": file_path, "page": i + 1}))
    return docs


def _load_docx(file_path: str) -> List[Document]:
    """加载 Word (.docx) 文件，含段落和表格"""
    from docx import Document as DocxDoc
    doc = DocxDoc(file_path)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text and c.text.strip():
                    parts.append(c.text)
    if not parts:
        return []
    return [Document(page_content="\n".join(parts), metadata={"source": file_path})]


_LOADER_MAP = {".txt": _load_txt, ".md": _load_md,
               ".pdf": _load_pdf, ".docx": _load_docx}


def _load_single(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    return _LOADER_MAP.get(ext, _load_txt)(file_path)


# ── 知识库类 ───────────────────────────────────────────────

class LocalKnowledgeBase:

    def __init__(
        self,
        model_name: str = "deepseek-r1:8b",
        embedding_model: str = "nomic-embed-text",
        persist_dir: str = "./chroma_db",
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        self.model_name = model_name
        self.persist_dir = persist_dir

        print(f"[*] Embedding: {embedding_model}")
        self.embeddings = OllamaHTTPEmbeddings(embedding_model, base_url=OLLAMA_BASE_URL)

        print(f"[*] LLM: {model_name}")
        self.llm = OllamaLLM(
            model=model_name, base_url=OLLAMA_BASE_URL,
            temperature=0.1, num_ctx=6144, num_predict=768,
            top_p=0.9, repeat_penalty=1.1,
        )

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )

        self.vector_store: Optional[Chroma] = None

        # 附注判别 Prompt（精简版，减少 token 加速推理）
        self.judge_prompt = PromptTemplate(
            template=(
                "你是工程图纸附注审查专家。基于规范上下文，判断附注是否合规。\n\n"
                "规范上下文：\n{context}\n\n"
                "附注：{question}\n\n"
                "回答格式（三行，简洁）：\n"
                "结果：【符合规范/不符合规范/规范中无明确规定】\n"
                "依据：引用条文简述理由\n"
                "条文：摘录关键规范原文"
            ),
            input_variables=["context", "question"],
        )

    @staticmethod
    def clean_answer(text: str) -> str:
        """去掉 DeepSeek R1 的 <think> 标签，保留后续内容"""
        cleaned = re.sub(r"<think>[\s\S]*?</think>", "", text).strip()
        # 如果 think 外为空，尝试从 think 内部提取结论（兜底）
        if not cleaned:
            inner = re.search(r"<think>[\s\S]*?</think>", text)
            if inner:
                # 取 think 最后几行（通常包含结论）
                lines = inner.group().strip().split("\n")
                cleaned = "\n".join(lines[-5:])
        return cleaned

    # ── 文档加载 ──────────────────────────────────────────

    def load_documents(self, docs_path: str) -> List[Document]:
        path = Path(docs_path)
        documents = []
        if path.is_file():
            try:
                documents = _load_single(str(path))
            except Exception as e:
                raise RuntimeError(
                    f"加载文件失败: {path.name}\n"
                    f"类型: {type(e).__name__}\n"
                    f"详情: {e}"
                ) from e
        else:
            import glob as g
            for ext, pat in {".txt": "**/*.txt", ".md": "**/*.md",
                             ".pdf": "**/*.pdf", ".docx": "**/*.docx"}.items():
                for fp in g.glob(str(path / pat), recursive=True):
                    try:
                        documents.extend(_load_single(fp))
                        print(f"  [OK] {os.path.basename(fp)}")
                    except Exception as e:
                        print(f"  [SKIP] {os.path.basename(fp)}: {e}")
        print(f"Total documents: {len(documents)}")
        return documents

    # ── 构建知识库 ────────────────────────────────────────

    def build_knowledge_base(self, docs_path: str):
        print(f"\n[*] Building KB from: {docs_path}")
        documents = self.load_documents(docs_path)
        if not documents:
            print("  [INFO] No documents found — initializing empty KB")
            return self.initialize_empty()
        print("[*] Splitting...")
        chunks = self.text_splitter.split_documents(documents)
        print(f"  [OK] {len(chunks)} chunks")
        print("[*] Embedding & storing...")
        self.vector_store = Chroma.from_documents(
            documents=chunks, embedding=self.embeddings,
            persist_directory=self.persist_dir, collection_name="knowledge_base",
        )
        print(f"  [OK] Saved to: {self.persist_dir}")
        return self

    def initialize_empty(self):
        """初始化空向量库（后续可 add_documents 追加）"""
        self.vector_store = Chroma(
            persist_directory=self.persist_dir,
            embedding_function=self.embeddings,
            collection_name="knowledge_base",
        )
        print(f"  [OK] Empty KB initialized at: {self.persist_dir}")
        return self

    # ── 加载已有知识库 ────────────────────────────────────

    def load_existing_kb(self):
        if not os.path.exists(self.persist_dir):
            raise FileNotFoundError(f"KB not found: {self.persist_dir}")
        self.vector_store = Chroma(
            persist_directory=self.persist_dir,
            embedding_function=self.embeddings,
            collection_name="knowledge_base",
        )
        print("[OK] KB loaded")
        return self

    # ── 添加文档（增量更新） ──────────────────────────────

    def add_documents(self, docs_path: str):
        if not self.vector_store:
            raise RuntimeError("Initialize KB first")
        documents = self.load_documents(docs_path)
        if not documents:
            raise ValueError(
                f"未能从文件中提取任何文本: {os.path.basename(docs_path)}\n"
                f"可能原因: 扫描版 PDF（图片型）需先用 OCR 工具提取文字，"
                f"或该文件本身不含文字内容"
            )
        chunks = self.text_splitter.split_documents(documents)
        if not chunks:
            raise ValueError(f"文本分割后无有效块: {os.path.basename(docs_path)}")
        self.vector_store.add_documents(chunks)
        print(f"[OK] Added {len(chunks)} chunks")
        return len(chunks)

    # ── 相似度检索 ────────────────────────────────────────

    def similarity_search(self, query: str, k: int = 5):
        if not self.vector_store:
            raise RuntimeError("KB not initialized")
        return self.vector_store.similarity_search(query, k=k)

    def mmr_search(self, query: str, k: int = 4, fetch_k: int = 12,
                   lambda_mult: float = 0.5):
        """MMR 检索 — 平衡相关性与多样性，避免结果来自同一文档"""
        if not self.vector_store:
            raise RuntimeError("KB not initialized")
        return self.vector_store.max_marginal_relevance_search(
            query, k=k, fetch_k=fetch_k, lambda_mult=lambda_mult,
        )

    # ── 附注判别（核心方法）──────────────────────────────

    def judge_annotation(self, annotation_text: str, k: int = 4) -> dict:
        """对附注条目进行判别，返回结果、依据和相关条文"""
        if not self.vector_store:
            raise RuntimeError("KB not initialized")

        t0 = time.time()

        # Step 1: MMR 检索 — 保证多文档覆盖
        docs = self.mmr_search(annotation_text, k=k, fetch_k=12, lambda_mult=0.5)
        sources_seen = set()
        for d in docs:
            sources_seen.add(os.path.basename(d.metadata.get("source", "")))
        t1 = time.time()
        print(f"  [MMR] {k} chunks from {len(sources_seen)} docs ({t1 - t0:.1f}s)")

        context = "\n\n".join(
            f"[{os.path.basename(d.metadata.get('source', ''))}]\n{d.page_content}"
            for i, d in enumerate(docs)
        )

        # Step 2: 用 DeepSeek R1 判别
        prompt_text = self.judge_prompt.format(
            context=context, question=annotation_text
        )

        raw = self.llm.invoke(prompt_text)
        t2 = time.time()

        # 兜底：LangChain invoke 偶尔返回空，直接用 HTTP API
        if not raw:
            try:
                resp = requests.post(
                    f"{OLLAMA_BASE_URL}/api/generate",
                    json={"model": self.model_name, "prompt": prompt_text,
                          "stream": False,
                          "options": {"temperature": 0.1, "num_ctx": 6144}},
                    timeout=120,
                )
                raw = resp.json().get("response", "")
            except Exception:
                pass

        print(f"  [LLM] inference {t2 - t1:.1f}s, output {len(raw)} chars")
        cleaned = self.clean_answer(raw)

        # Step 3: 解析结果（宽松匹配，兼容多种格式）
        result_type = "规范中无明确规定"
        basis = ""
        clause = ""

        # 匹配 "结果：符合规范" / "判断结果：【符合规范】" 等变体
        m = re.search(
            r"(?:判断)?结果[：:]\s*【?\s*(符合规范|不符合规范|规范中无明确规定)\s*】?",
            cleaned,
        )
        if m:
            result_type = m.group(1)

        # 匹配 "依据：..." / "判断依据：..."
        m = re.search(
            r"(?:判断)?依据[：:]\s*(.+?)(?=(?:相关规范)?条文[：:]|$)",
            cleaned, re.DOTALL,
        )
        if m:
            basis = m.group(1).strip()

        m = re.search(r"(?:相关规范)?条文[：:]\s*(.+)", cleaned, re.DOTALL)
        if m:
            clause = m.group(1).strip()

        return {
            "annotation": annotation_text,
            "result": result_type,
            "basis": basis or cleaned,
            "clause": clause,
            "raw_answer": cleaned,
            "sources": [
                {"source": os.path.basename(d.metadata.get("source", "")),
                 "page": d.metadata.get("page", ""),
                 "snippet": d.page_content[:200]}
                for d in docs
            ],
        }

    # ── 通用问答 ──────────────────────────────────────────

    def ask(self, question: str, k: int = 4) -> dict:
        if not self.vector_store:
            raise RuntimeError("KB not initialized")
        qa_prompt = PromptTemplate(
            template=(
                "你是一个专业技术规范助手。基于以下上下文回答问题。\n"
                "如果上下文中无相关信息，请明确告知。\n\n"
                "上下文：\n{context}\n\n"
                "问题：{question}\n\n回答："
            ),
            input_variables=["context", "question"],
        )
        from langchain_classic.chains import RetrievalQA
        qa = RetrievalQA.from_chain_type(
            llm=self.llm, chain_type="stuff",
            retriever=self.vector_store.as_retriever(
                search_type="similarity", search_kwargs={"k": k}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": qa_prompt},
        )
        result = qa.invoke({"query": question})
        return {
            "question": question,
            "answer": self.clean_answer(result["result"]),
            "sources": [
                {"source": os.path.basename(d.metadata.get("source", "")),
                 "snippet": d.page_content[:300]}
                for d in result["source_documents"]
            ],
        }

    # ── 列出知识库文档统计 ────────────────────────────────

    def get_stats(self) -> dict:
        if not self.vector_store:
            return {"status": "未初始化", "chunks": 0}
        try:
            collection = self.vector_store._collection
            return {
                "status": "已加载",
                "chunks": collection.count(),
                "persist_dir": self.persist_dir,
            }
        except Exception:
            return {"status": "已加载", "chunks": "?", "persist_dir": self.persist_dir}
